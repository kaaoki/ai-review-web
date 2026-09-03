"""
review_engine.py
文書の抽出とGemini APIによる自動レビューのロジック。
元のCLI版（review_tool.py）のレビュー観点・出力形式を踏襲している。
"""

import io
import json
import os

from google import genai

# ---- 対応フォーマット ----
SUPPORTED_EXTENSIONS = [".txt", ".md", ".docx", ".xlsx", ".xlsm", ".pdf"]

# 現時点(2026年9月)でのGemini APIの標準的なFlashモデル。
# モデルは頻繁に更新されるため、Google AI Studioで最新のモデルIDを
# 確認し、必要であれば環境変数 GEMINI_MODEL で上書きしてください。
DEFAULT_MODEL = os.environ.get("GEMINI_MODEL", "gemini-3.8-flash")

REVIEW_PROMPT_TEMPLATE = """あなたは日本語の業務文書をチェックする校閲者です。
以下の文書を読み、次の4つの観点で問題点を洗い出してください。

1. 誤字脱字
2. 日付・数値などの事実関係の誤り（文書内で矛盾する日付や数値など）
3. 文章内での矛盾（前半と後半で方針や結論が食い違っている箇所など）
4. フォーマット・表記ゆれ

一見矛盾しているように見えても、論理的に両立しうる記述（例：件数増加と売上減少の併記など）は指摘しないでください。

必ず次のJSON形式のみで出力してください。前置きや説明文、Markdownのコードブロック記法は一切不要です。

{{
  "issues": [
    {{
      "location": "該当箇所（見出しや行の内容を短く引用）",
      "content": "指摘内容",
      "suggestion": "修正提案",
      "severity": "高|中|低"
    }}
  ]
}}

問題が見つからない場合は "issues": [] としてください。

--- 文書本文 ---
{document_text}
"""


def extract_text(filename: str, file_bytes: bytes) -> str:
    """アップロードされたファイルからテキストを抽出する"""
    ext = os.path.splitext(filename)[1].lower()

    if ext in (".txt", ".md"):
        return file_bytes.decode("utf-8", errors="replace")

    if ext == ".docx":
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)

    if ext in (".xlsx", ".xlsm"):
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        parts = []
        for sheet in wb.worksheets:
            parts.append(f"[シート: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if ext == ".pdf":
        import pdfplumber

        parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
        return "\n".join(parts)

    raise ValueError(f"未対応のファイル形式です: {ext}")


def review_document(document_text: str, api_key: str, model: str = DEFAULT_MODEL) -> dict:
    """Gemini APIで文書をレビューし、構造化された結果を返す"""
    if not document_text.strip():
        return {"issues": []}

    client = genai.Client(api_key=api_key)
    prompt = REVIEW_PROMPT_TEMPLATE.format(document_text=document_text[:60000])

    response = client.models.generate_content(
        model=model,
        contents=prompt,
        config={
            "response_mime_type": "application/json",
        },
    )

    raw_text = response.text.strip()
    # コードブロック記法が付いてしまった場合の保険
    if raw_text.startswith("```"):
        raw_text = raw_text.strip("`")
        raw_text = raw_text.replace("json\n", "", 1)

    try:
        result = json.loads(raw_text)
    except json.JSONDecodeError:
        result = {"issues": [], "parse_error": True, "raw_response": raw_text}

    if "issues" not in result:
        result["issues"] = []

    return result


def summarize_result(result: dict) -> str:
    """一覧表示用の短いサマリー文字列を作る"""
    issues = result.get("issues", [])
    if not issues:
        return "指摘なし"
    counts = {"高": 0, "中": 0, "低": 0}
    for issue in issues:
        sev = issue.get("severity", "低")
        if sev in counts:
            counts[sev] += 1
    return f"計{len(issues)}件（高:{counts['高']} 中:{counts['中']} 低:{counts['低']}）"
